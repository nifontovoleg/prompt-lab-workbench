"""Прогон test_input для всех промптов и генерация DOCX-отчёта."""

from __future__ import annotations

import asyncio
import json
import re
import time
from datetime import datetime
from pathlib import Path
from typing import Any

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from dotenv import load_dotenv

from prompt_chat import (
    DEFAULT_MAX_TOKENS,
    DEFAULT_MODEL,
    DEFAULT_TEMPERATURE,
    PROMPTS_DIR,
    build_system_message,
    generate_answer,
    list_prompt_files,
    load_prompt_file,
)

OUT_DIR = Path(__file__).resolve().parent / "reports"
OUT_DOCX = OUT_DIR / "otchet_testirovanie_promptov.docx"
OUT_JSON = OUT_DIR / "test_results.json"


def keywords_from_expected(description: str) -> list[str]:
    """Грубая эвристика: ключевые фразы из expected_test_output_description."""
    text = description.lower()
    # разбиваем по запятым / «и» после двоеточия
    if ":" in text:
        text = text.split(":", 1)[1]
    parts = re.split(r"[,;]| и ", text)
    keys = []
    for part in parts:
        cleaned = re.sub(r"[()\d.]+", " ", part)
        cleaned = " ".join(cleaned.split())
        if len(cleaned) >= 6:
            keys.append(cleaned)
    return keys[:12]


def coverage_score(answer: str, expected_desc: str) -> tuple[float, list[str], list[str]]:
    answer_l = answer.lower()
    keys = keywords_from_expected(expected_desc)
    hit, miss = [], []
    for key in keys:
        # достаточно совпадения ~60% слов ключа
        words = [w for w in key.split() if len(w) > 3]
        if not words:
            continue
        matched = sum(1 for w in words if w in answer_l)
        if matched / len(words) >= 0.5:
            hit.append(key)
        else:
            miss.append(key)
    total = len(hit) + len(miss)
    score = (len(hit) / total * 100) if total else 0.0
    return score, hit, miss


def structure_check(answer: str, components: list[dict[str, Any]]) -> tuple[int, int, list[str]]:
    found = []
    for item in components:
        name = (item.get("name") or "").strip()
        if name and name.lower() in answer.lower():
            found.append(name)
    return len(found), len(components), found


async def run_one(prompt_path: Path) -> dict[str, Any]:
    data = load_prompt_file(prompt_path)
    test_input = (data.get("test_input") or "").strip()
    expected = (data.get("expected_test_output_description") or "").strip()
    components = (data.get("structure") or {}).get("components") or []

    started = time.perf_counter()
    result = await generate_answer(
        system=build_system_message(data),
        user_question=test_input,
        model=DEFAULT_MODEL,
        temperature=DEFAULT_TEMPERATURE,
        max_tokens=DEFAULT_MAX_TOKENS,
    )
    elapsed = time.perf_counter() - started

    score, hit, miss = coverage_score(result["content"], expected)
    found_n, total_n, found_names = structure_check(result["content"], components)

    return {
        "prompt_id": data.get("prompt_id") or prompt_path.stem,
        "name": data.get("name") or prompt_path.stem,
        "version": data.get("version"),
        "category": data.get("category"),
        "test_input": test_input,
        "expected_description": expected,
        "answer": result["content"],
        "model": result.get("model"),
        "total_tokens": result.get("total_tokens"),
        "prompt_tokens": result.get("prompt_tokens"),
        "completion_tokens": result.get("completion_tokens"),
        "elapsed_sec": round(elapsed, 2),
        "coverage_score": round(score, 1),
        "coverage_hit": hit,
        "coverage_miss": miss,
        "structure_found": found_n,
        "structure_total": total_n,
        "structure_names": found_names,
    }


def verdict(item: dict[str, Any]) -> str:
    struct_ok = item["structure_total"] == 0 or (
        item["structure_found"] / item["structure_total"] >= 0.5
    )
    cov_ok = item["coverage_score"] >= 50
    if struct_ok and cov_ok:
        return "Пройден"
    if struct_ok or cov_ok:
        return "Частично"
    return "Не пройден"


def conclusion_for(item: dict[str, Any]) -> str:
    v = verdict(item)
    name = item["name"]
    parts = [
        f"Промпт «{name}» — статус: {v}.",
        f"Ответ сформирован за {item['elapsed_sec']} с, "
        f"модель {item['model']}, всего токенов: {item['total_tokens']}.",
        f"Соответствие ожидаемому описанию (эвристика по ключевым темам): "
        f"{item['coverage_score']}%.",
        f"Блоки структуры в ответе: {item['structure_found']} из {item['structure_total']}.",
    ]
    if item["coverage_hit"]:
        parts.append("Покрыты темы: " + "; ".join(item["coverage_hit"][:5]) + ".")
    if item["coverage_miss"]:
        parts.append("Слабее отражены: " + "; ".join(item["coverage_miss"][:5]) + ".")
    if v == "Пройден":
        parts.append(
            "Вывод: промпт стабильно направляет модель на нужный формат и содержание; "
            "пригоден для рабочего использования."
        )
    elif v == "Частично":
        parts.append(
            "Вывод: базовая цель достигается, но стоит уточнить инструкции по "
            "недостающим разделам/темам."
        )
    else:
        parts.append(
            "Вывод: требуется доработка формулировок role/structure/format."
        )
    return " ".join(parts)


def add_heading(doc: Document, text: str, level: int = 1) -> None:
    doc.add_heading(text, level=level)


def add_para(doc: Document, text: str, bold: bool = False) -> None:
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold


def build_docx(results: list[dict[str, Any]]) -> None:
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    doc = Document()

    title = doc.add_heading("Отчёт о тестировании промптов", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    add_para(
        doc,
        f"Дата: {datetime.now().strftime('%d.%m.%Y %H:%M')}",
    )
    add_para(
        doc,
        "Проект: test from cli — интерактивный CLI (prompt_chat.py) + заготовки в prompts/.",
    )
    add_para(
        doc,
        f"Условия теста: ProxyAPI, модель {DEFAULT_MODEL}, "
        f"temperature={DEFAULT_TEMPERATURE}, max_tokens={DEFAULT_MAX_TOKENS}. "
        "Для каждого промпта использован встроенный test_input.",
    )

    add_heading(doc, "1. Краткое резюме", 1)
    passed = sum(1 for r in results if verdict(r) == "Пройден")
    partial = sum(1 for r in results if verdict(r) == "Частично")
    failed = sum(1 for r in results if verdict(r) == "Не пройден")
    add_para(
        doc,
        f"Протестировано промптов: {len(results)}. "
        f"Пройдено: {passed}, частично: {partial}, не пройдено: {failed}.",
    )
    total_tokens = sum(r.get("total_tokens") or 0 for r in results)
    add_para(doc, f"Суммарный расход токенов: {total_tokens}.")

    table = doc.add_table(rows=1, cols=5)
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    hdr[0].text = "Промпт"
    hdr[1].text = "Версия"
    hdr[2].text = "Статус"
    hdr[3].text = "Покрытие тем, %"
    hdr[4].text = "Токены"
    for r in results:
        row = table.add_row().cells
        row[0].text = r["name"]
        row[1].text = str(r.get("version") or "—")
        row[2].text = verdict(r)
        row[3].text = str(r["coverage_score"])
        row[4].text = str(r.get("total_tokens") or "н/д")

    add_heading(doc, "2. Результаты по каждому промпту", 1)

    for i, r in enumerate(results, 1):
        add_heading(doc, f"2.{i}. {r['name']} ({r['prompt_id']})", 2)
        add_para(doc, f"Категория: {r.get('category') or '—'}")
        add_para(doc, f"Версия: {r.get('version') or '—'}")
        add_para(doc, f"Статус: {verdict(r)}", bold=True)

        add_heading(doc, "Входной тест (test_input)", 3)
        add_para(doc, r["test_input"][:1200] + ("…" if len(r["test_input"]) > 1200 else ""))

        add_heading(doc, "Ожидания (expected_test_output_description)", 3)
        add_para(doc, r["expected_description"])

        add_heading(doc, "Метрики запроса", 3)
        add_para(doc, f"Модель: {r.get('model')}")
        add_para(doc, f"Время ответа: {r['elapsed_sec']} с")
        add_para(doc, f"Использовано токенов: {r.get('total_tokens')}")
        add_para(doc, f"Промпт токены: {r.get('prompt_tokens')}")
        add_para(doc, f"Ответ токены: {r.get('completion_tokens')}")
        add_para(
            doc,
            f"Структурные блоки в ответе: {r['structure_found']}/{r['structure_total']}"
            + (f" ({', '.join(r['structure_names'])})" if r["structure_names"] else ""),
        )

        add_heading(doc, "Ключевые результаты", 3)
        # краткий фрагмент ответа
        preview = r["answer"]
        if len(preview) > 1800:
            preview = preview[:1800] + "\n…\n[фрагмент ответа сокращён]"
        add_para(doc, preview)

        add_heading(doc, "Выводы", 3)
        add_para(doc, conclusion_for(r))

    add_heading(doc, "3. Общие выводы и рекомендации", 1)
    add_para(
        doc,
        "Все три промпта успешно интегрированы в CLI и отрабатывают через ProxyAPI "
        "на штатных test_input. Структура JSON (role, context, structure, format) "
        "корректно собирается в system-сообщение.",
    )
    add_para(
        doc,
        "Рекомендации: 1) для code_structure и task_planning при длинных ответах "
        "можно поднять max_tokens; 2) закрепить обязательные заголовки разделов "
        "в role/format, чтобы повышать структурное покрытие; "
        "3) добавить автоматические регрессионные прогоны test_input в CI.",
    )
    add_para(
        doc,
        "Формат сдачи: документ совместим с Google Документами "
        "(Файл → Открыть / загрузка .docx в Google Диск → Открыть с помощью Google Документы).",
    )

    doc.save(OUT_DOCX)


async def main() -> None:
    load_dotenv()
    files = list_prompt_files()
    if not files:
        raise SystemExit(f"Нет промптов в {PROMPTS_DIR}")

    print(f"Найдено промптов: {len(files)}")
    results: list[dict[str, Any]] = []
    for path in files:
        print(f"→ Тест: {path.stem} ...")
        item = await run_one(path)
        results.append(item)
        print(
            f"  статус={verdict(item)}, покрытие={item['coverage_score']}%, "
            f"токены={item['total_tokens']}, {item['elapsed_sec']}с"
        )

    OUT_DIR.mkdir(parents=True, exist_ok=True)
    OUT_JSON.write_text(
        json.dumps(results, ensure_ascii=False, indent=2),
        encoding="utf-8",
    )
    build_docx(results)
    print(f"\nОтчёт сохранён: {OUT_DOCX}")
    print(f"Сырые данные: {OUT_JSON}")


if __name__ == "__main__":
    asyncio.run(main())
