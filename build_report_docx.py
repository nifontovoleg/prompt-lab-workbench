"""Красивый DOCX-отчёт в стиле рабочей тетради интенсива."""

from __future__ import annotations

import json
from datetime import datetime
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

RESULTS_PATH = Path("reports/test_results.json")
OUT = Path("reports/otchet_testirovanie_promptov.docx")

# Палитра в духе учебной тетради
BLUE = RGBColor(0x1A, 0x56, 0xDB)
DARK = RGBColor(0x1F, 0x29, 0x37)
GRAY = RGBColor(0x4B, 0x55, 0x63)
GREEN = RGBColor(0x05, 0x7A, 0x55)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
HEADER_BG = "1A56DB"
OK_BG = "D1FAE5"
SOFT_BG = "EFF6FF"
ROW_ALT = "F8FAFC"


def set_run_font(run, *, size: int = 11, bold: bool = False, color: RGBColor = DARK, name: str = "Calibri") -> None:
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    run.font.size = Pt(size)
    run.bold = bold
    run.font.color.rgb = color


def set_cell_shading(cell, hex_color: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), hex_color)
    shd.set(qn("w:val"), "clear")
    tc_pr.append(shd)


def set_cell_text(
    cell,
    text: str,
    *,
    bold: bool = False,
    color: RGBColor = DARK,
    size: int = 10,
    align=WD_ALIGN_PARAGRAPH.LEFT,
) -> None:
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = align
    run = p.add_run(text)
    set_run_font(run, size=size, bold=bold, color=color)


def add_title(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(text)
    set_run_font(run, size=26, bold=True, color=BLUE)


def add_subtitle(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(14)
    run = p.add_run(text)
    set_run_font(run, size=12, color=GRAY)


def add_heading_emoji(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(16)
    p.paragraph_format.space_after = Pt(8)
    run = p.add_run(text)
    set_run_font(run, size=14, bold=True, color=BLUE)


def add_body(doc: Document, text: str, *, bold: bool = False, color: RGBColor = DARK) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    run = p.add_run(text)
    set_run_font(run, size=11, bold=bold, color=color)


def add_bullet(doc: Document, text: str) -> None:
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(3)
    # очищаем дефолтный текст стиля и пишем свой
    if p.runs:
        p.runs[0].text = text
        set_run_font(p.runs[0], size=11, color=DARK)
    else:
        run = p.add_run(text)
        set_run_font(run, size=11, color=DARK)


def add_separator(doc: Document) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(8)
    run = p.add_run("—" * 28)
    set_run_font(run, size=10, color=RGBColor(0xCB, 0xD5, 0xE1))


def add_callout(doc: Document, title: str, body: str) -> None:
    table = doc.add_table(rows=1, cols=1)
    table.autofit = True
    cell = table.rows[0].cells[0]
    set_cell_shading(cell, SOFT_BG)
    cell.text = ""
    p1 = cell.paragraphs[0]
    r1 = p1.add_run(title)
    set_run_font(r1, size=11, bold=True, color=BLUE)
    p2 = cell.add_paragraph()
    r2 = p2.add_run(body)
    set_run_font(r2, size=10, color=DARK)
    doc.add_paragraph()


def style_header_row(row, headers: list[str]) -> None:
    for i, text in enumerate(headers):
        set_cell_text(row.cells[i], text, bold=True, color=WHITE, size=10, align=WD_ALIGN_PARAGRAPH.CENTER)
        set_cell_shading(row.cells[i], HEADER_BG)


def main() -> None:
    results = json.loads(RESULTS_PATH.read_text(encoding="utf-8"))
    by_id = {r["prompt_id"]: r for r in results}

    def get(pid: str) -> dict:
        return by_id[pid]

    order = ["summary", "code_structure", "task_planning"]

    conclusions = {
        "summary": (
            "Промпт стабильно делает краткое структурированное резюме: главная идея, "
            "ключевые пункты, факты и выводы. Для 100% формальной полноты можно явно "
            "требовать отдельный блок «Заголовок»."
        ),
        "code_structure": (
            "Промпт даёт полноценную архитектурную карту e-commerce (FastAPI + React). "
            "Все 8 структурных блоков присутствуют. Для длинных ответов лучше поднять max_tokens."
        ),
        "task_planning": (
            "Промпт формирует рабочий план MVP на 8 недель с декомпозицией, приоритетами, "
            "рисками и чекпоинтами. Улучшение — явная понедельная таблица с ответственными."
        ),
    }

    key_points = {
        "summary": [
            "Структура ответа: основная идея → ключевые пункты → важные детали → выводы.",
            "Сохранены факты: инвестиции, этика/приватность, горизонт 5–10 лет, баланс инноваций и ответственности.",
            "Деловой стиль без домыслов; задача краткого изложения выполнена.",
        ],
        "code_structure": [
            "Покрыты все 8 блоков: архитектура, директории, модули, паттерны, API, БД, конфигурация, тесты.",
            "Учтены JWT, PostgreSQL, заказы, админка, email, платежи, отзывы, хранилище изображений.",
            "Есть дерево backend/frontend и перечень REST endpoint’ов.",
        ],
        "task_planning": [
            "Все 9 блоков плана присутствуют, включая зависимости, риски и чекпоинты.",
            "Учтены ограничения команды (2 FE + 1 BE + дизайнер) и дедлайн 8 недель.",
            "Есть приоритизация Must / Should / Nice to Have.",
        ],
    }

    doc = Document()
    section = doc.sections[0]
    section.top_margin = Cm(1.8)
    section.bottom_margin = Cm(1.8)
    section.left_margin = Cm(2.0)
    section.right_margin = Cm(2.0)

    # —— Шапка в стиле тетради ——
    add_title(doc, "Рабочая тетрадь")
    add_subtitle(doc, "Отчёт о тестировании промптов · Интенсив по промпт-инжинирингу 🚀")

    add_body(
        doc,
        "Привет! Это итоговый отчёт по практическому заданию: протестировать три промпта из проекта, "
        "зафиксировать ключевые результаты и сформулировать выводы по каждому.",
    )
    add_body(doc, "Используй этот документ, чтобы:")
    for t in [
        "✅ Показать прогресс по промптам и статусы выполнения.",
        "✅ Зафиксировать метрики запросов (модель, токены, время).",
        "✅ Сравнить ожидаемый результат с фактическим ответом модели.",
        "✅ Сформулировать выводы и рекомендации по доработке.",
    ]:
        add_bullet(doc, t)

    add_separator(doc)

    # —— Информация о проекте ——
    add_heading_emoji(doc, "📌 Информация о проекте")
    info = doc.add_table(rows=3, cols=2)
    info.style = "Table Grid"
    rows_info = [
        ("📦 Проект", "prompt-lab-workbench"),
        ("📅 Дата отчёта", datetime.now().strftime("%d.%m.%Y")),
        ("🛠 Стек", "Python · prompt_chat.py · ProxyAPI · prompts/"),
    ]
    for i, (k, v) in enumerate(rows_info):
        set_cell_text(info.rows[i].cells[0], k, bold=True, size=10, color=BLUE)
        set_cell_shading(info.rows[i].cells[0], SOFT_BG)
        set_cell_text(info.rows[i].cells[1], v, size=10)

    doc.add_paragraph()
    add_body(doc, "💡 Цели прогона:", bold=True)
    for t in [
        "Цель 1 — проверить набор промптов (summary / code_structure / task_planning).",
        "Цель 2 — убедиться, что CLI корректно собирает system prompt и ходит в ProxyAPI.",
        "Цель 3 — зафиксировать метрики, ключевые результаты и выводы по каждому промпту.",
    ]:
        add_bullet(doc, t)

    add_separator(doc)

    # —— План / статусы ——
    add_heading_emoji(doc, "🗂 План тестирования промптов")
    plan = doc.add_table(rows=1, cols=4)
    plan.style = "Table Grid"
    style_header_row(plan.rows[0], ["Задача", "Промпт", "Статус", "Заметки"])
    plan_rows = [
        (
            "Тест резюмирования текста",
            "summary · v1.1",
            "Готово ✅",
            f"{get('summary')['total_tokens']} ток. · {get('summary')['elapsed_sec']} с",
        ),
        (
            "Тест генерации структуры кода",
            "code_structure · v1.7",
            "Готово ✅",
            f"{get('code_structure')['total_tokens']} ток. · {get('code_structure')['elapsed_sec']} с",
        ),
        (
            "Тест планирования задач",
            "task_planning · v1.4",
            "Готово ✅",
            f"{get('task_planning')['total_tokens']} ток. · {get('task_planning')['elapsed_sec']} с",
        ),
    ]
    for i, (task, prompt, status, note) in enumerate(plan_rows):
        row = plan.add_row()
        set_cell_text(row.cells[0], task, size=10)
        set_cell_text(row.cells[1], prompt, size=10)
        set_cell_text(row.cells[2], status, bold=True, color=GREEN, size=10, align=WD_ALIGN_PARAGRAPH.CENTER)
        set_cell_shading(row.cells[2], OK_BG)
        set_cell_text(row.cells[3], note, size=10)
        if i % 2 == 1:
            for c in (0, 1, 3):
                set_cell_shading(row.cells[c], ROW_ALT)

    doc.add_paragraph()
    total_tokens = sum((x.get("total_tokens") or 0) for x in results)
    total_time = sum(x["elapsed_sec"] for x in results)
    add_callout(
        doc,
        "📊 Сводка прогона",
        f"Условия: gpt-4o-mini · temperature=0.7 · max_tokens=2000 · ProxyAPI. "
        f"Протестировано промптов: 3/3. Суммарно токенов: {total_tokens}. "
        f"Суммарное время: {total_time:.1f} с. Все тесты — статус «Готово».",
    )

    add_separator(doc)

    # —— Детали по каждому промпту ——
    add_heading_emoji(doc, "🧪 Результаты тестирования")

    titles = {
        "summary": "🎯 Промпт 1. Резюме текста",
        "code_structure": "🏗 Промпт 2. Генерация структуры кода",
        "task_planning": "🗓 Промпт 3. Планирование задач",
    }

    for pid in order:
        item = get(pid)
        add_heading_emoji(doc, titles[pid])
        add_body(
            doc,
            f"Категория: {item.get('category')} · Версия: {item.get('version')} · Статус: Готово ✅",
            color=GRAY,
        )

        add_body(doc, "📝 Тестовый вход (test_input):", bold=True)
        ti = item["test_input"]
        add_body(doc, ti if len(ti) <= 900 else ti[:900] + "…")

        add_body(doc, "✅ Ожидаемый результат:", bold=True)
        add_body(doc, item["expected_description"])

        add_body(doc, "📈 Метрики запроса:", bold=True)
        metrics = doc.add_table(rows=1, cols=2)
        metrics.style = "Table Grid"
        style_header_row(metrics.rows[0], ["Параметр", "Значение"])
        metric_rows = [
            ("Модель", str(item.get("model"))),
            ("Время ответа", f"{item['elapsed_sec']} с"),
            ("Использовано токенов", str(item.get("total_tokens"))),
            ("Промпт токены", str(item.get("prompt_tokens"))),
            ("Ответ токены", str(item.get("completion_tokens"))),
            (
                "Структурные блоки",
                f"{item['structure_found']} из {item['structure_total']}",
            ),
        ]
        for j, (k, v) in enumerate(metric_rows):
            row = metrics.add_row()
            set_cell_text(row.cells[0], k, bold=True, size=10, color=BLUE)
            set_cell_shading(row.cells[0], SOFT_BG)
            set_cell_text(row.cells[1], v, size=10)
            if j % 2 == 1:
                set_cell_shading(row.cells[1], ROW_ALT)

        doc.add_paragraph()
        add_body(doc, "🔑 Ключевые результаты:", bold=True)
        for bullet in key_points[pid]:
            add_bullet(doc, bullet)

        add_body(doc, "📤 Фрагмент ответа модели:", bold=True)
        preview = item["answer"]
        if len(preview) > 1100:
            preview = preview[:1100] + "\n…\n[фрагмент сокращён]"
        add_callout(doc, "Ответ модели", preview)

        add_body(doc, "💡 Вывод:", bold=True)
        add_body(doc, conclusions[pid])
        add_separator(doc)

    # —— Итог ——
    add_heading_emoji(doc, "🚀 Общие выводы и рекомендации")
    for t in [
        "Интеграция prompts/ → system prompt → ProxyAPI работает корректно для всех трёх сценариев.",
        "Самый быстрый и «лёгкий» — summary; code_structure и task_planning дают длинные ответы.",
        "Структура components соблюдается: code_structure и task_planning — полностью, summary — почти полностью.",
        "Рекомендации: поднять max_tokens для длинных промптов; явно требовать секцию «Заголовок» в summary; "
        "оставить автопрогон через run_prompt_tests.py.",
    ]:
        add_bullet(doc, t)

    doc.add_paragraph()
    add_callout(
        doc,
        "📎 Формат сдачи",
        "Документ оформлен как рабочая тетрадь и готов к загрузке в Google Документы: "
        "Google Диск → загрузить .docx → «Открыть с помощью Google Документы». "
        f"Файл: {OUT.resolve()}",
    )

    footer = doc.add_paragraph()
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fr = footer.add_run("Готов(а) к следующему этапу интенсива? Тогда продолжаем! 🚀")
    set_run_font(fr, size=11, bold=True, color=BLUE)

    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT)
    print(f"Updated: {OUT.resolve()}")


if __name__ == "__main__":
    main()
